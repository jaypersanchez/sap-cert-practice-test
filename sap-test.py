import argparse
import tkinter as tk
from tkinter import messagebox
from docx import Document
import random
import openai
import os
import pickle
from dotenv import load_dotenv

# Load the OpenAI API key from the .env file
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Function to parse questions from the Practice Test Document
def parse_questions(doc_path):
    doc = Document(doc_path)
    questions = []
    question = {}
    
    # Define keywords for tagging
    tag_keywords = {
        'SAP BTP': ['BTP', 'Business Technology Platform'],
        'Cloud Foundry': ['Cloud Foundry', 'cf push'],
        'SAP CAP': ['CAP', 'Core Data Services', 'CDS'],
        'Integration': ['Integration', 'API', 'Integration Suite'],
        'Side-by-Side Extensions': ['Side-by-Side', 'extension'],
        'Data Analytics': ['Analytics', 'SAP HANA'],
        'Automation': ['automation', 'process automation']
    }
    
    for para in doc.paragraphs:
        if para.text.startswith("Q"):
            if question and question['options']:  # Only append if there are options
                # Automatically assign tags based on keywords
                question['tags'] = assign_tags(question['question'], tag_keywords)
                questions.append(question)
            # Reset for new question
            question = {'question': para.text.replace("?", ""), 'options': [], 'answer': '', 'explanation': '', 'type': 'standard'}
        elif para.text.startswith("A.") or para.text.startswith("B.") or para.text.startswith("C.") or para.text.startswith("D."):
            if question:  # Only add options if a question is being built
                question['options'].append(para.text)
        elif para.text.startswith("Answer:"):
            if question:  # Only set answer if a question is being built
                question['answer'] = para.text.split(":")[1].strip().lower()  # Store answer in lowercase for consistency
        elif para.text.startswith("Referenced from"):
            if question:  # Only set explanation if a question is being built
                question['explanation'] = para.text
    
    # Add the last question if it has options
    if question and question['options']:
        question['tags'] = assign_tags(question['question'], tag_keywords)
        questions.append(question)
    
    return questions

# Helper function to assign tags based on keywords
def assign_tags(question_text, tag_keywords):
    tags = []
    for tag, keywords in tag_keywords.items():
        if any(keyword.lower() in question_text.lower() for keyword in keywords):
            tags.append(tag)
    return tags

# Function to load study references from the condensed notes document
def load_references_from_notes(notes_path):
    doc = Document(notes_path)
    references = []
    for para in doc.paragraphs:
        if para.text:
            references.append(para.text)
    return references


# Load question history from file or return an empty set
def load_question_history():
    if os.path.exists("question_history.pkl"):
        with open("question_history.pkl", "rb") as f:
            return pickle.load(f)
    return set()

# Save question history to file
def save_question_history(history):
    with open("question_history.pkl", "wb") as f:
        pickle.dump(history, f)

# Get rotating questions to avoid repetition
def get_rotating_questions(questions, num=10, num_sets=5):
    history = load_question_history()

    # Divide questions into subsets for rotation
    subset_size = len(questions) // num_sets
    subsets = [questions[i * subset_size:(i + 1) * subset_size] for i in range(num_sets)]

    # Select a subset that hasn't been fully used
    for subset in subsets:
        if all(q['question'] not in history for q in subset):
            # Ensure we sample only as many questions as are available
            selected_questions = random.sample(subset, min(num, len(subset)))
            break
    else:
        # If all subsets have been used, reset history
        history.clear()
        # Ensure we sample only as many questions as are available
        selected_questions = random.sample(subsets[0], min(num, len(subsets[0])))

    # Update history
    history.update(q['question'] for q in selected_questions)
    save_question_history(history)

    return selected_questions


# Get questions based on tags (categories) with handling for smaller sample sizes
def get_tag_based_random_questions(questions, num=10):
    # Define categories by tags
    categorized_questions = {}
    
    for question in questions:
        for tag in question.get('tags', []):
            if tag not in categorized_questions:
                categorized_questions[tag] = []
            categorized_questions[tag].append(question)

    selected_questions = []
    
    # Ensure we sample from each category
    for tag, q_list in categorized_questions.items():
        if q_list:
            selected_questions.extend(random.sample(q_list, min(2, len(q_list))))

    # Shuffle and return only as many questions as needed
    random.shuffle(selected_questions)
    return selected_questions[:min(num, len(selected_questions))]



# Main GUI Application
class PracticeTestApp:
    def __init__(self, master, questions, references, mode="rotating"):
        self.master = master
        self.master.title("SAP Certification Practice Test")

        # Switch between rotating or tag-based strategy
        if mode == "rotating":
            self.questions = get_rotating_questions(questions, num=10)
        elif mode == "tag":
            self.questions = get_tag_based_random_questions(questions, num=10)

        self.references = references
        self.current_question = 0
        self.user_answers = []

        # Question Label
        self.question_label = tk.Label(master, text="", wraplength=400, justify="left")
        self.question_label.pack(pady=20)

        # Options Radio Buttons
        self.options_var = tk.StringVar()
        self.option_buttons = []
        for i in range(4):  # Assuming 4 options (A, B, C, D)
            btn = tk.Radiobutton(master, text="", variable=self.options_var, value=i, wraplength=400, anchor='w', justify="left")
            btn.pack(anchor='w', pady=5)
            self.option_buttons.append(btn)

        # Next Button
        self.next_button = tk.Button(master, text="Next", command=self.next_question)
        self.next_button.pack(pady=20)

        self.load_question()

    def load_question(self):
        # Load current question into the interface
        q = self.questions[self.current_question]
        self.question_label.config(text=q['question'])

        for i, option in enumerate(q['options']):
            self.option_buttons[i].config(text=option)
        self.options_var.set(None)  # Reset selection

    def next_question(self):
        if self.options_var.get() is None:
            messagebox.showwarning("No selection", "Please select an answer before proceeding.")
            return

        # Save user's answer
        self.user_answers.append(self.options_var.get())
        self.current_question += 1

        # Check if the test is complete
        if self.current_question < len(self.questions):
            self.load_question()
        else:
            self.show_results()  # Ensure this method is defined in the class

    def get_chatgpt_explanation(self, result_summary):
        """Function to send the result summary to OpenAI API and get a response with explanations."""
        prompt = (
            f"{result_summary}\n\n"
            "Can you provide detailed explanations for each answer? "
            "Especially explain the correct answers and clarify the mistakes in the wrong answers."
        )
        
        try:
            # Using the ChatCompletion API with gpt-3.5-turbo or gpt-4
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # You can switch to "gpt-4" if needed
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that provides detailed explanations."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,  # Adjust token count based on needs
                temperature=0.7
            )
            return response['choices'][0]['message']['content'].strip()
        except Exception as e:
            return f"Error: {str(e)}"

    # Method to show results and fetch detailed explanations from ChatGPT
    def show_results(self):
        correct = 0
        results_window = tk.Toplevel(self.master)
        results_window.title("Results")

        # Create a Text widget to display results on the left side (60% width)
        result_text_widget = tk.Text(results_window, wrap="word", width=50, height=20)
        result_text_widget.grid(row=0, column=0, padx=10, pady=20)

        # Create another Text widget for ChatGPT explanations on the right side (40% width)
        chatgpt_response_widget = tk.Text(results_window, wrap="word", width=40, height=20)
        chatgpt_response_widget.grid(row=0, column=1, padx=10, pady=20)

        # Configure tags to color text
        result_text_widget.tag_config('wrong', foreground="red")
        result_text_widget.tag_config('correct', foreground="green")

        # Evaluate answers and build the result summary
        result_summary = "Quiz Results:\n"
        for i, q in enumerate(self.questions):
            selected = self.user_answers[i]
            correct_answer = q['answer']  # The correct answer is a string, e.g., 'b'
            question_text = q['question']  # Store the original question

            # Determine the index of the correct option
            correct_option = None
            for j, option in enumerate(q['options']):
                if option.lower().startswith(correct_answer):
                    correct_option = j
                    break

            selected_option_text = q['options'][int(selected)]  # Get the full text of the selected option
            correct_option_text = q['options'][correct_option]  # Get the full text of the correct option

            # Add to result summary and insert into the Text widget
            result_text_widget.insert(tk.END, f"Q{i + 1}. {question_text}\n")
            if int(selected) == correct_option:
                correct += 1
                # Show correct answers in green
                result_text_widget.insert(tk.END, f"Your Answer: {selected_option_text} (Correct)\n\n", 'correct')
                result_summary += f"Q{i + 1}: Correct\nYour Answer: {selected_option_text}\n"
            else:
                # Show wrong answers in red
                result_text_widget.insert(tk.END, f"Your Answer: {selected_option_text} (Wrong)\n", 'wrong')
                result_text_widget.insert(tk.END, f"Correct Answer: {correct_option_text} (Correct)\n\n", 'correct')
                result_summary += f"Q{i + 1}: Wrong\nYour Answer: {selected_option_text}\nCorrect Answer: {correct_option_text}\n"

        # Show the final score
        score_text = f"Your Score: {correct}/{len(self.questions)}\n\n"
        result_text_widget.insert(tk.END, score_text)

        # Add the score to the result summary
        result_summary += score_text

        # Make the Text widget read-only
        result_text_widget.config(state=tk.DISABLED)

        # Fetch the explanation from ChatGPT using the test results summary
        chatgpt_explanation = self.get_chatgpt_explanation(result_summary)

        # Insert the response into the right-side Text widget
        chatgpt_response_widget.insert(tk.END, chatgpt_explanation)
        chatgpt_response_widget.config(state=tk.DISABLED)

        # OK Button to close the results window
        ok_button = tk.Button(results_window, text="OK", command=results_window.destroy)
        ok_button.grid(row=1, column=0, columnspan=2, pady=10)


# Main function to start the application
def main():
    # Use argparse to get the mode from the command-line arguments
    parser = argparse.ArgumentParser(description="SAP Certification Practice Test")
    parser.add_argument('--mode', choices=['rotating', 'tag'], default='rotating', help="Choose the mode for the test: rotating or tag.")
    args = parser.parse_args()

    # Parse the questions and references
    questions = parse_questions("PracticeTestChatGptGenerated.docx")
    references = load_references_from_notes("CondensedStudyNotesforCertification.docx")

    # Start the GUI
    root = tk.Tk()
    app = PracticeTestApp(root, questions, references, mode=args.mode)
    root.mainloop()


if __name__ == "__main__":
    main()
