import sqlite3
from docx import Document
from PyPDF2 import PdfReader
import os

def create_database(db_name):
    """Create a SQLite database and tables for questions and references if they do not exist."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    # Create table for questions
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject TEXT NOT NULL,
            question TEXT NOT NULL,
            options TEXT NOT NULL,
            answer TEXT NOT NULL,
            explanation TEXT,
            tags TEXT
        )
    ''')

    # Create table for study references
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS study_references (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject TEXT NOT NULL,
            reference TEXT NOT NULL
        )
    ''')
    
    conn.commit()
    conn.close()

def insert_question(db_name, subject, question, options, answer, explanation='', tags=''):
    """Insert a question into the database."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO questions (subject, question, options, answer, explanation, tags)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (subject, question, options, answer, explanation, tags))
    
    conn.commit()
    conn.close()

def parse_and_insert_questions(db_name, subject, doc_path):
    """Parse questions from DOCX or PDF and insert them into the database."""
    if doc_path.endswith('.pdf'):
        questions = parse_questions_from_pdf(doc_path)
    elif doc_path.endswith('.docx'):
        questions = parse_questions_from_docx(doc_path)
    else:
        raise ValueError("Unsupported file format. Please provide a DOCX or PDF file.")
    
    for q in questions:
        options = ', '.join(q['options'])  # Join options into a single string
        insert_question(db_name, subject, q['question'], options, q['answer'], q['explanation'], ', '.join(q['tags']))

def parse_questions_from_docx(doc_path):
    """Parse questions from a DOCX file."""
    questions = []
    doc = Document(doc_path)
    question = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Q"):  # Check for question
            if question and question.get('options'):
                questions.append(question)
            question = {
                'question': text.replace("?", ""), 
                'options': [], 
                'answer': '', 
                'explanation': '', 
                'tags': []
            }
        elif any(text.startswith(opt) for opt in ["A.", "B.", "C.", "D."]):
            if question:
                question['options'].append(text)
        elif text.startswith("Answer:"):
            if question:
                question['answer'] = text.split(":")[1].strip().lower()  # Store answer in lowercase
        elif text.startswith("Referenced from"):
            if question:
                question['explanation'] = text

    # Append the last question if available
    if question and question.get('options'):
        questions.append(question)
    
    return questions

def parse_questions_from_pdf(doc_path):
    """Parse questions from a PDF file."""
    questions = []
    reader = PdfReader(doc_path)
    question = {}
    
    for page in reader.pages:
        text = page.extract_text()
        for line in text.splitlines():
            line = line.strip()
            if line.startswith("Q"):  # Check for question
                if question and question.get('options'):
                    questions.append(question)
                question = {
                    'question': line.replace("?", "").strip(),  # Ensure to clean up the question text
                    'options': [], 
                    'answer': '', 
                    'explanation': '', 
                    'tags': []
                }
            elif any(line.startswith(opt) for opt in ["A.", "B.", "C.", "D."]):
                if question:
                    question['options'].append(line.strip())  # Clean up option text
            elif line.startswith("Answer:"):
                if question:
                    question['answer'] = line.split(":")[1].strip().lower()  # Store answer in lowercase
            elif line.startswith("Referenced from"):
                if question:
                    question['explanation'] = line.strip()

    # Append the last question if available
    if question and question.get('options'):
        questions.append(question)
    
    return questions

def insert_questions_from_subfolders(base_dir, db_name):
    # Create the database and table
    create_database(db_name)

    # Walk through all subdirectories in the base directory
    for root, dirs, files in os.walk(base_dir):
        for file in files:
            file_path = os.path.join(root, file)
            subject = os.path.basename(root)  # Get the name of the current folder as the subject
            
            # Check the file extension and parse accordingly
            if file.endswith('.docx'):
                print(f"Inserting questions from {file_path} for subject '{subject}'")
                parse_and_insert_questions(db_name, subject, file_path)
            elif file.endswith('.pdf'):
                print(f"Inserting questions from {file_path} for subject '{subject}'")
                parse_and_insert_questions(db_name, subject, file_path)

def insert_reference(db_name, subject, reference):
    """Insert a reference into the database."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO study_references (subject, reference)
        VALUES (?, ?)
    ''', (subject, reference))
    
    conn.commit()
    conn.close()

def load_references_from_notes(db_name, subject, notes_path):
    """Load study references from a DOCX or PDF file containing notes and insert into the database."""
    if notes_path.endswith('.docx'):
        references = parse_references_from_docx(notes_path)
    elif notes_path.endswith('.pdf'):
        references = parse_references_from_pdf(notes_path)
    else:
        raise ValueError("Unsupported file format. Please provide a DOCX or PDF file.")
    
    for reference in references:
        insert_reference(db_name, subject, reference)

def parse_references_from_docx(doc_path):
    """Parse references from a DOCX file."""
    references = []
    doc = Document(doc_path)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Only add non-empty references
            references.append(text)
    
    return references

def parse_references_from_pdf(doc_path):
    """Parse references from a PDF file."""
    references = []
    reader = PdfReader(doc_path)
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            references.extend(text.splitlines())  # Split lines and add to references
    
    return references

def get_references_by_subject(db_name, subject):
    """Retrieve and display references for a specific subject."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute("SELECT reference FROM study_references WHERE subject = ?", (subject,))
    rows = cursor.fetchall()
    
    if rows:
        print(f"References for Subject: {subject}")
        for row in rows:
            print(f"- {row[0]}")
    else:
        print(f"No references found for subject: {subject}")
    
    conn.close()

if __name__ == "__main__":
    base_directory = 'contents'  # Base directory containing subfolders
    db_name = 'questions.db'  # Database name
    insert_questions_from_subfolders(base_directory, db_name)
    print("All questions have been successfully inserted into the database.")
