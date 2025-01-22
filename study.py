import argparse
import tkinter as tk
from tkinter import messagebox, simpledialog
from tkinter import ttk
from docx import Document
import random
import os
import pickle
import requests
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import sqlite3

# Load the OpenAI API key from the .env file
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Define tag keywords for tagging questions
tag_keywords = {
    'Math': ['math', 'algebra', 'geometry'],
    'Science': ['science', 'biology', 'chemistry'],
    'History': ['history', 'historical', 'ancient'],
    # Add more tags and their corresponding keywords as needed
}

# ----------------------------
# Document and Question Parsing
# ----------------------------

def parse_questions(doc_path):
    """
    Parse questions from a DOCX or PDF file.
    This function assumes the questions are tagged with specific markers.
    """
    if doc_path.endswith('.pdf'):
        return parse_questions_from_pdf(doc_path)
    elif doc_path.endswith('.docx'):
        return parse_questions_from_docx(doc_path)  # Assuming you have a function to parse DOCX
    else:
        raise ValueError("Unsupported file format. Please provide a DOCX or PDF file.")

def parse_questions_from_docx(doc_path):
    """
    Parse questions from a DOCX file.
    This function assumes the questions are tagged with specific markers.
    """
    questions = []
    doc = Document(doc_path)
    question = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Q"):  # Check for question
            if question and question.get('options'):
                question['tags'] = assign_tags(question['question'], tag_keywords)
                questions.append(question)
            question = {
                'question': text.replace("?", ""), 
                'options': [], 
                'answer': '', 
                'explanation': '', 
                'type': 'standard'
            }
        elif any(text.startswith(opt) for opt in ["A.", "B.", "C.", "D."]):
            if question:
                question['options'].append(text)
        elif text.startswith("Answer:"):
            if question:
                question['answer'] = text.split(":")[1].strip().lower()  # Store answer in lowercase
        elif text.startswith("Referenced from"):
            if question:
                question['explanation'] = text

    # Append the last question if available
    if question and question.get('options'):
        question['tags'] = assign_tags(question['question'], tag_keywords)
        questions.append(question)
    
    return questions

def parse_questions_from_pdf(doc_path):
    """
    Parse questions from a PDF file.
    This function assumes the questions are tagged with specific markers.
    """
    questions = []
    reader = PdfReader(doc_path)
    question = {}
    
    for page in reader.pages:
        text = page.extract_text()
        for line in text.splitlines():
            line = line.strip()
            print(f"Parsing line: {line}")  # Debugging line
            if line.startswith("Q"):  # Check for question
                if question and question.get('options'):
                    question['tags'] = assign_tags(question['question'], tag_keywords)
                    questions.append(question)
                question = {
                    'question': line.replace("?", "").strip(),  # Ensure to clean up the question text
                    'options': [], 
                    'answer': '', 
                    'explanation': '', 
                    'type': 'standard'
                }
            elif any(line.startswith(opt) for opt in ["A.", "B.", "C.", "D."]):
                if question:
                    question['options'].append(line.strip())  # Clean up option text
            elif line.startswith("Answer:"):
                if question:
                    question['answer'] = line.split(":")[1].strip().lower()  # Store answer in lowercase
            elif line.startswith("Referenced from"):
                if question:
                    question['explanation'] = line.strip()

    # Append the last question if available
    if question and question.get('options'):
        question['tags'] = assign_tags(question['question'], tag_keywords)
        questions.append(question)
    
    return questions


def assign_tags(question_text, tag_keywords):
    """Assign tags to a question based on provided keywords."""
    tags = []
    for tag, keywords in tag_keywords.items():
        if any(keyword.lower() in question_text.lower() for keyword in keywords):
            tags.append(tag)
    return tags


def load_references_from_notes(notes_path):
    """Load study references from a DOCX or PDF file containing notes."""
    references = []
    
    if notes_path.endswith('.docx'):
        doc = Document(notes_path)
        for para in doc.paragraphs:
            if para.text:
                references.append(para.text)
    elif notes_path.endswith('.pdf'):
        references = load_references_from_pdf(notes_path)  # New function to handle PDF
    else:
        raise ValueError("Unsupported file format. Please provide a DOCX or PDF file.")
    
    return references

def load_references_from_pdf(notes_path):
    """Load study references from a PDF file containing notes."""
    references = []
    reader = PdfReader(notes_path)
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            references.extend(text.splitlines())  # Split lines and add to references
    
    return references


# ----------------------------
# History Management for Questions
# ----------------------------

def load_question_history():
    if os.path.exists("question_history.pkl"):
        with open("question_history.pkl", "rb") as f:
            return pickle.load(f)
    return set()


def save_question_history(history):
    with open("question_history.pkl", "wb") as f:
        pickle.dump(history, f)


def get_rotating_questions(questions, num=10, num_sets=5):
    history = load_question_history()

    # Divide questions into subsets for rotation
    subset_size = len(questions) // num_sets
    subsets = [questions[i * subset_size:(i + 1) * subset_size] for i in range(num_sets)]

    # Select a subset that hasn't been fully used
    for subset in subsets:
        if all(q['question'] not in history for q in subset):
            selected_questions = random.sample(subset, min(num, len(subset)))
            break
    else:
        # All subsets used; reset history
        history.clear()
        selected_questions = random.sample(subsets[0], min(num, len(subsets[0])))

    # Update history
    history.update(q['question'] for q in selected_questions)
    save_question_history(history)
    return selected_questions


def get_tag_based_random_questions(questions, num=10):
    categorized_questions = {}
    for question in questions:
        for tag in question.get('tags', []):
            categorized_questions.setdefault(tag, []).append(question)

    selected_questions = []
    for tag, q_list in categorized_questions.items():
        if q_list:
            selected_questions.extend(random.sample(q_list, min(2, len(q_list))))
    random.shuffle(selected_questions)
    return selected_questions[:min(num, len(selected_questions))]


# ----------------------------
# Study Topic Selection Helpers
# ----------------------------

def select_topic(content_dir="contents"):
    """
    Create a simple window for topic selection.
    This function lists all sub-folders in content_dir and returns the selected topic.
    """
    # Get list of topics (each sub-folder in content_dir)
    topics = [d for d in os.listdir(content_dir) if os.path.isdir(os.path.join(content_dir, d))]
    if not topics:
        messagebox.showerror("Error", f"No topics found in the '{content_dir}' folder.")
        exit(1)
    
    # Create a temporary Tk window for selection
    sel_window = tk.Toplevel()
    sel_window.title("Select Study Topic")
    tk.Label(sel_window, text="Select a study topic:").pack(pady=10)

    # Variable to store the selection
    topic_var = tk.StringVar(value=topics[0])
    
    for topic in topics:
        tk.Radiobutton(sel_window, text=topic, variable=topic_var, value=topic).pack(anchor='w', padx=20)

    def confirm_selection():
        sel_window.destroy()
    
    tk.Button(sel_window, text="Select", command=confirm_selection).pack(pady=10)
    # Wait for the user to close the selection window
    sel_window.wait_window()
    return topic_var.get()


def load_topic_documents(topic_folder):
    """
    In the given topic folder, search for a question file and a notes file.
    The file names should include 'question' (or 'questions') and 'note' (or 'notes'),
    respectively. Returns the full paths for (question_file, notes_file).
    """
    question_file = None
    notes_file = None
    for filename in os.listdir(topic_folder):
        lower = filename.lower()
        if "question" in lower or "questions" in lower:
            if filename.endswith('.docx') or filename.endswith('.pdf'):  # Check for both formats
                question_file = os.path.join(topic_folder, filename)
        elif "note" in lower or "notes" in lower:
            if filename.endswith('.docx') or filename.endswith('.pdf'):  # Check for both formats
                notes_file = os.path.join(topic_folder, filename)
    return question_file, notes_file


# ----------------------------
# Database Functions
# ----------------------------

def load_subjects(db_name):
    """Load all distinct subjects from the questions table."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute("SELECT DISTINCT subject FROM questions")
    subjects = cursor.fetchall()
    
    conn.close()
    return [subject[0] for subject in subjects]  # Extracting the subject names

def load_questions(db_name):
    """Load all questions from the database."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM questions")
    questions = cursor.fetchall()
    
    conn.close()
    return questions

def load_references(db_name, subject):
    """Load references for a specific subject."""
    conn = sqlite3.connect(db_name)
    cursor = conn.cursor()
    
    cursor.execute("SELECT reference FROM study_references WHERE subject = ?", (subject,))
    references = cursor.fetchall()
    
    conn.close()
    return [ref[0] for ref in references]  # Extracting the reference text


# ----------------------------
# Main GUI Application
# ----------------------------

class SubjectSelectionApp:
    def __init__(self, master, db_name):
        self.master = master
        self.master.title("Select Subject")
        
        self.db_name = db_name
        self.subjects = load_subjects(db_name)
        
        self.label = tk.Label(master, text="Select a subject for the test:")
        self.label.pack(pady=10)
        
        self.subject_var = tk.StringVar()
        self.subject_combobox = ttk.Combobox(master, textvariable=self.subject_var, values=self.subjects)
        self.subject_combobox.pack(pady=10)
        
        self.test_type_var = tk.StringVar(value="Preview")  # Default to Preview
        self.test_type_label = tk.Label(master, text="Select test type:")
        self.test_type_label.pack(pady=10)
        
        self.preview_radio = tk.Radiobutton(master, text="Preview (10 questions)", variable=self.test_type_var, value="Preview")
        self.preview_radio.pack(anchor='w')
        
        self.practice_radio = tk.Radiobutton(master, text="Practice (25 questions)", variable=self.test_type_var, value="Practice")
        self.practice_radio.pack(anchor='w')
        
        self.certified_radio = tk.Radiobutton(master, text="Certified Simulation (All questions)", variable=self.test_type_var, value="Certified")
        self.certified_radio.pack(anchor='w')
        
        self.select_button = tk.Button(master, text="Start Test", command=self.start_test)
        self.select_button.pack(pady=20)
    
    def start_test(self):
        selected_subject = self.subject_var.get()
        if not selected_subject:
            messagebox.showwarning("No Subject Selected", "Please select a subject.")
            return
        
        questions = load_questions(self.db_name)  # Load all questions
        references = load_references(self.db_name, selected_subject)  # Load references for the selected subject
        
        # Determine the number of questions based on the selected test type
        test_type = self.test_type_var.get()
        if test_type == "Preview":
            questions = random.sample(questions, min(10, len(questions)))  # 10 random questions
        elif test_type == "Practice":
            questions = random.sample(questions, min(25, len(questions)))  # 25 random questions
        elif test_type == "Certified":
            questions = questions  # All questions
        
        # Open the test window without destroying the main window
        self.open_test_window(questions, references)
    
    def open_test_window(self, questions, references):
        test_window = tk.Toplevel(self.master)
        app = PracticeTestApp(test_window, questions, references)

class PracticeTestApp:
    def __init__(self, master, questions, references):
        self.master = master
        self.master.title("Practice Test")
        
        self.questions = questions
        self.references = references
        self.current_question = 0
        self.user_answers = []
        
        # GUI Widgets
        self.question_label = tk.Label(master, text="", wraplength=800, justify="left")
        self.question_label.pack(pady=20)
        
        self.question_count_label = tk.Label(master, text="", wraplength=800, justify="left")
        self.question_count_label.pack(pady=5)

        self.options_var = tk.StringVar()
        self.option_buttons = []
        
        self.next_button = tk.Button(master, text="Next", command=self.next_question)
        self.next_button.pack(pady=20)
        
        self.load_question()
    
    def load_question(self):
        if self.current_question >= len(self.questions):
            messagebox.showerror("Error", "No more questions available.")
            return
        q = self.questions[self.current_question]
        self.question_label.config(text=q[2])  # Question text
        options = q[3].split(', ')  # Assuming options are stored as a comma-separated string
        
        # Clear existing option buttons
        for btn in self.option_buttons:
            btn.pack_forget()  # Remove the button from the window

        # Create new option buttons based on the number of options
        self.option_buttons = []  # Reset the option buttons list
        for i, option in enumerate(options):
            btn = tk.Radiobutton(self.master, text=option, variable=self.options_var, value=i, wraplength=800, anchor='w', justify="left")
            btn.pack(anchor='w', pady=5)
            self.option_buttons.append(btn)

        # Reset the selected option to None
        self.options_var.set(None)  # Clear any previous selection
        
        # Update the question count label
        self.question_count_label.config(text=f"Question {self.current_question + 1} of {len(self.questions)}")
    
    def next_question(self):
        if self.options_var.get() is None:
            messagebox.showwarning("No selection", "Please select an answer before proceeding.")
            return
        self.user_answers.append(self.options_var.get())
        self.current_question += 1
        if self.current_question < len(self.questions):
            self.load_question()
        else:
            self.show_results()
    
    def show_results(self):
        correct = 0
        results_window = tk.Toplevel(self.master)
        results_window.title("Results")
        
        result_text_widget = tk.Text(results_window, wrap="word", width=50, height=20)
        result_text_widget.grid(row=0, column=0, padx=10, pady=20)
        
        result_summary = "Quiz Results:\n"
        for i, q in enumerate(self.questions):
            selected = self.user_answers[i]
            correct_answer = q[4]  # Assuming the correct answer is stored in the 5th column
            question_text = q[2]
            
            selected_option_text = q[3].split(', ')[int(selected)]  # Get the selected option text
            correct_option_text = correct_answer  # Assuming the correct answer is stored as a single letter
            
            result_text_widget.insert(tk.END, f"Q{i + 1}. {question_text}\n")
            if selected_option_text.lower() == correct_option_text.lower():
                correct += 1
                result_text_widget.insert(tk.END, f"Your Answer: {selected_option_text} (Correct)\n\n")
                result_summary += f"Q{i + 1}: Correct\nYour Answer: {selected_option_text}\n"
            else:
                result_text_widget.insert(tk.END, f"Your Answer: {selected_option_text} (Wrong)\n")
                result_text_widget.insert(tk.END, f"Correct Answer: {correct_option_text} (Correct)\n\n")
                result_summary += f"Q{i + 1}: Wrong\nYour Answer: {selected_option_text}\nCorrect Answer: {correct_option_text}\n"
        
        score_text = f"Your Score: {correct}/{len(self.questions)}\n\n"
        result_text_widget.insert(tk.END, score_text)
        result_summary += score_text
        result_text_widget.config(state=tk.DISABLED)
        
        # Show references for the subject
        references_text = "\nReferences:\n" + "\n".join(self.references)
        result_text_widget.insert(tk.END, references_text)
        
        ok_button = tk.Button(results_window, text="OK", command=results_window.destroy)
        ok_button.grid(row=1, column=0, columnspan=2, pady=10)


# ----------------------------
# Main Function
# ----------------------------

def main():
    db_name = 'questions.db'  # Database name
    root = tk.Tk()
    app = SubjectSelectionApp(root, db_name)
    root.mainloop()


if __name__ == "__main__":
    main()
